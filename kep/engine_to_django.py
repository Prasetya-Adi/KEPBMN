# Create your views here.
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from django.db.models import Sum

from .models import Barang
from django.db.models import Sum, Count
from django.db.models import F

from docx import shared, Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
#import locale


def make_bdn_2(id_skep, no_kep_bdn, tgl_kep_bdn, no_sbp, tgl_sbp):
    # Get Data
    all = Barang.objects.filter(nomor_skep=id_skep).values_list(
        'nomor_skep__no_bdn',
        'nomor_skep__no_bmn',
        'jenis_barang__jenis_barang',
        'merek',
        'isi',
        'jenis_barang__isi_satuan',
        'jumlah_kemasan',
        'jenis_barang__jumlah_kemasan_satuan',
        'harga_jual_eceran',
        'jenis_barang__jumlah_kemasan_satuan',
        'nilai',)

    def set_col_widths_record(table):
        widths = (shared.Cm(0.8),
                  shared.Cm(2.5),
                  shared.Cm(0.8),
                  shared.Cm(1.75),
                  shared.Cm(0.8),
                  shared.Cm(1.75),
                  shared.Cm(0.8),
                  shared.Cm(1.75),
                  shared.Cm(2),
                  )
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    def make_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'KEP'
                    for run in paragraph.runs:
                        run.font.bold = True

    def make_rows_font(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'KEP'

    def make_table_record(document):
        table = document.add_table(rows=1, cols=9)
        #table.style = 'table'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Merek'
        hdr_cells[2].text = 'Isi'
        #hdr_cells[3].text = ''
        hdr_cells[4].text = 'Jumlah'
        #hdr_cells[5].text = ''
        hdr_cells[6].text = 'Total'
        #hdr_cells[7].text = ''
        hdr_cells[8].text = 'HJE'

        # HORIZONTAL ALIGNMENT
        for i in range(9):
            cell = table.cell(0, i)
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # vertical aligment
        for i in range(9):
            table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # MAKE BOLD
        make_rows_bold(table.rows[0])

        # MERGE CELL
        a = table.cell(0, 2)
        b = table.cell(0, 3)
        c = table.cell(0, 4)
        d = table.cell(0, 5)
        e = table.cell(0, 6)
        f = table.cell(0, 7)

        A = a.merge(b)
        B = c.merge(d)
        B = e.merge(f)

        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        return table

    def make_table_header(document, jenis, jumlah_satuan_jenis):
        table = document.add_table(rows=9, cols=4)
        # ROW 1
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Menimbang'
        hdr_cells[1].text = ':'
        hdr_cells[2].text = 'a.'
        hdr_cells[3].text = 'Bahwa di Kantor Pengawasan dan Pelayanan Bea dan Cukai Tipe Madya Pabean C Tasikmalaya terdapat '+jumlah_satuan_jenis + \
            ' tanpa dilekati pita cukai karena adanya pelanggaran terhadap Undang-Undang Nomor 11 Tahun 1995 tentang Cukai sebagaimana telah diubah dengan Undang-Undang Nomor 39 Tahun 2007.'

        # ROW 2
        hdr_cells = table.rows[1].cells
        hdr_cells[2].text = 'b.'
        hdr_cells[3].text = 'Bahwa sehubungan dengan hal tersebut diatas, dipandang perlu untuk menetapkan barang barang tersebut sebagai Barang yang Dikuasai Negara.'

        # ROW 3
        hdr_cells = table.rows[2].cells

        # ROW 4
        hdr_cells = table.rows[3].cells
        hdr_cells[0].text = 'Mengingat'
        hdr_cells[1].text = ':'
        hdr_cells[2].text = '1.'
        hdr_cells[3].text = 'Undang-Undang Nomor 11 tahun 1995  tentang Cukai (Lembaran  Negara Republik Indonesia Tahun 1995 Nomor 76, Tambahan Lembaran Negara Republik Indonesia Nomor 3613) sebagaimana telah diubah dengan  Undang-Undang  Nomor 39 tahun 2007 (Lembaran Negara Republik Indonesia Tahun 2007 Nomor 105, Tambahan Lembaran Negara Republik Indonesia Nomor 4755);'

        # ROW 5
        hdr_cells = table.rows[4].cells

        hdr_cells[2].text = '2.'
        hdr_cells[3].text = 'Peraturan Menteri Keuangan Republik Indonesia Nomor: 39/PMK.04/2014 tanggal 19 Februari 2014 tentang Tata Cara Penyelesaian Barang Kena Cukai dan Barang - barang Lain yang Dirampas Untuk Negara atau yang Dikuasai Negara;'

        # ROW 6
        hdr_cells = table.rows[5].cells
        hdr_cells[2].text = '3.'
        hdr_cells[3].text = 'Peraturan Direktur Jenderal Bea dan Cukai nomor PER-17/BC/2020 tanggal 30 Desember 2020 tentang Tata Laksana Pengawasan.'

        # ROW 7
        hdr_cells = table.rows[6].cells

        # ROW 8
        hdr_cells = table.rows[7].cells
        hdr_cells[3].text = 'MEMUTUSKAN\n'
        table.cell(7, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ROW 9
        hdr_cells = table.rows[8].cells
        hdr_cells[0].text = 'Menetapkan'
        hdr_cells[1].text = ':'
        hdr_cells[3].text = 'KEPUTUSAN KEPALA KANTOR PENGAWASAN DAN PELAYANAN  BEA DAN CUKAI TIPE MADYA PABEAN C TASIKMALAYA TENTANG PENETAPAN BARANG KENA CUKAI BERUPA ' + \
            jenis.upper()+' SEBAGAI BARANG DIKUASAI NEGARA'

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.style = 'KEP-1'
        widths = (shared.Cm(2.52),
                  shared.Cm(0.49),
                  shared.Cm(0.49),
                  shared.Cm(13.7),
                  )
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        return table

    #locale.setlocale(locale.LC_ALL, 'id_ID')
    # NO SBP
    no_sbp = "SBP-"+no_sbp+"/WBC.09/KPP.MP.06/"+tgl_sbp.strftime('%Y')
    tgl_sbp = tgl_sbp.strftime('%-d %B %Y')

    # KEP BDN
    no_kep_bdn = "KEP-"+str(no_kep_bdn) + \
        "/WBC.09/KPP.MP.06/"+tgl_kep_bdn.strftime('%Y')
    tgl_kep_bdn = tgl_kep_bdn.strftime('%-d %B %Y')

    # JUMLAH JENIS BARANG
    # jumlah = Barang.objects.filter(
    #     nomor_skep=id_skep).aggregate(Sum('jumlah_kemasan'))
    # jum = jumlah['jumlah_kemasan__sum']
    # jum_sat = all[0][7]
    # jenis = all[0][2]

    # jumlah_satuan_jenis = str(jum)+' '+str(jum_sat)+' '+str(jenis)

    # JUMLAH JENIS BARANG CARA BARU

    barang = Barang.objects.filter(nomor_skep=id_skep)

    def join_koma_dan(l):
        if not l:
            return ""
        elif len(l) == 1:
            return l[0]
        else:
            return ', '.join(l[:-1]) + " dan " + l[-1]

    barang_annotate = barang.values(jb=F('jenis_barang__jenis_barang'), sat=F('jenis_barang__jumlah_kemasan_satuan')).annotate(
        jumlah=Sum('jumlah_kemasan'))
    # JENIS BARANG TEXT
    list_barang = []
    for b in barang_annotate:
        text = str(b['jb'])
        list_barang.append(text)

    # FULL TEXT
    list_full = []
    for b in barang_annotate:
        text = str(b['jumlah'])+' ' + \
            str(b['sat'])+' '+str(b['jb'])
        list_full.append(text)

    jenis = join_koma_dan(list_barang)
    jumlah_satuan_jenis = join_koma_dan(list_full)

    # PENETAPAN
    lokasi = 'Tasikmalaya'
    kepala_kantor = 'Indriya Karyadi'
    nip = '19710905 199201 1 004'

    # SET DOCUMENT FORMAT
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(330)
    section.page_width = Mm(215)
    section.top_margin = Mm(20.0)
    section.bottom_margin = Mm(15.0)
    section.left_margin = Mm(25.0)
    section.right_margin = Mm(20.0)
    section.header_distance = Mm(10.0)
    section.footer_distance = Mm(12.7)

    # BASIC STYLE
    styles = document.styles
    style = styles.add_style('KEP', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    paragraph = style.paragraph_format
    font.name = 'Bookman Old Style'
    font.size = Pt(11)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.space_after = Pt(0)

    # BASIC STYLE AND JUSTIFY
    styles = document.styles
    style = styles.add_style('KEP-1', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    paragraph = style.paragraph_format
    font.name = 'Bookman Old Style'
    font.size = Pt(11)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.space_after = Pt(0)

    # STYLE INDENT AND JUSTIFY
    styles = document.styles
    style = styles.add_style('KEP-2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    paragraph = style.paragraph_format
    font.name = 'Bookman Old Style'
    font.size = Pt(11)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.space_after = Pt(0)
    paragraph.left_indent = Mm(32.5)

    # STYLE INDENT AND CENTER
    styles = document.styles
    style = styles.add_style('KEP-2-Center', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    paragraph = style.paragraph_format
    font.name = 'Bookman Old Style'
    font.size = Pt(11)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.space_after = Pt(0)
    paragraph.left_indent = Mm(32.5)

    para = document.add_paragraph(
        'KEMENTERIAN KEUANGAN REPUBLIK INDONESIA\nDIREKTORAT JENDERAL BEA DAN CUKAI\n\nKEPUTUSAN KEPALA KANTOR PENGAWASAN DAN PELAYANAN\nBEA DAN CUKAI TIPE MADYA PABEAN C TASIKMALAYA\n'+'NOMOR '+no_kep_bdn+'\n\nTENTANG\nPENETAPAN BARANG KENA CUKAI BERUPA\n'+jenis.upper()+'\nSEBAGAI BARANG DIKUASAI NEGARA\n\nKEPALA KANTOR PENGAWASAN DAN PELAYANAN BEA DAN CUKAI\nTIPE MADYA PABEAN C TASIKMALAYA\n\n', style='KEP')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #
    table_header = make_table_header(document, jenis, jumlah_satuan_jenis)

    para = document.add_paragraph('\nPasal 1', style='KEP-2-Center')
    para = document.add_paragraph('Barang-barang Kena Cukai berupa ' +
                                  jumlah_satuan_jenis+' dengan rincian:', style='KEP-2')

    table_record = make_table_record(document)

    # 'nomor_skep__no_bdn',
    # 'nomor_skep__no_bmn',
    # 'jenis_barang__jenis_barang',
    # 'merek',
    # 'isi',
    # 'jenis_barang__isi_satuan',
    # 'jumlah_kemasan',
    # 'jenis_barang__jumlah_kemasan_satuan',
    # 'harga_jual_eceran',
    # 'jenis_barang__jumlah_kemasan_satuan',
    # 'nilai',

    i = 1
    for bdn, bmn, jb, mer, isi, isisat, jm, jmsat, hjecer, hjesat, nilai in all:
        total = isi * jm
        row_cells = table_record.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(mer)
        row_cells[2].text = str(isi)
        row_cells[3].text = str(isisat)
        row_cells[4].text = str(jm)
        row_cells[5].text = str(jmsat)
        row_cells[6].text = str(total)
        row_cells[7].text = str(isisat)
        row_cells[8].text = 'Rp'+str(hjecer)+'\n/'+str(jmsat)
        make_rows_font(table_record.rows[i])
        i = i+1

    set_col_widths_record(table_record)

    para = document.add_paragraph(
        "adalah barang-barang hasil penindakan karena pelanggaran terhadap Undang-Undang Nomor 11 Tahun 1995 tentang Cukai sebagaimana telah diubah dengan Undang-Undang Nomor 39 tahun 2007, sesuai dengan Surat Bukti Penindakan nomor "+no_sbp+" tanggal "+tgl_sbp+".", style='KEP-2')

    para = document.add_paragraph('\nPasal 2', style='KEP-2-Center')
    para = document.add_paragraph(
        'Barang-barang sebagaimana dimaksud dalam Pasal 1 diatas sesuai dengan Pasal 54 Undang-Undang Nomor 11 Tahun 1995 tentang Cukai sebagaimana telah diubah dengan Undang-Undang Nomor 39 Tahun 2007, maka ditetapkan sebagai barang yang Dikuasai Negara.', style='KEP-2')

    para = document.add_paragraph('\nPasal 3', style='KEP-2-Center')
    para = document.add_paragraph(
        'Barang yang Dikuasai Negara sebagaimana dimaksud dalam Pasal 1 diatas ditimbun di Gudang Barang Hasil Penindakan dibawah Pengawasan Kantor Pengawasan dan Pelayanan Bea dan Cukai Tipe Madya Pabean C Tasikmalaya.', style='KEP-2')

    para = document.add_paragraph('\nPasal 4', style='KEP-2-Center')
    para = document.add_paragraph(
        'Keputusan ini mulai berlaku pada tanggal ditetapkan, dengan ketentuan apabila dikemudian hari terdapat kekeliruan akan diadakan perbaikan sebagaimana mestinya.', style='KEP-2')

    para = document.add_paragraph(
        'Salinan Keputusan Kepala Kantor Pengawasan dan Pelayanan Bea dan Cukai Tipe Madya Pabean C Tasikmalaya ini disampaikan kepada\n1. Direktur Teknis dan Fasilitas Cukai;\n2. Direktur Penindakan dan Penyidikan;\n3. Kepala Kantor Wilayah DJBC Jawa Barat.\n\n', style='KEP')
    para.paragraph_format.left_indent = Mm(32.5)

    para = document.add_paragraph(
        'Ditetapkan di : '+lokasi+'\nPada tanggal : '+tgl_kep_bdn+'\n\nKepala Kantor,\n\n\n\n'+kepala_kantor+'\nNIP '+nip, style='KEP')
    para.paragraph_format.left_indent = Mm(100.0)

    return document


def make_bmn_2(id_skep, no_kep_bmn, tgl_kep_bmn, no_kep_bdn, tgl_kep_bdn):
    all = Barang.objects.filter(nomor_skep=id_skep).values_list(
        'nomor_skep__no_bdn',
        'nomor_skep__no_bmn',
        'jenis_barang__jenis_barang',
        'merek',
        'isi',
        'jenis_barang__isi_satuan',
        'jumlah_kemasan',
        'jenis_barang__jumlah_kemasan_satuan',
        'harga_jual_eceran',
        'jenis_barang__jumlah_kemasan_satuan',
        'nilai',)

    def set_col_widths_record(table):
        widths = (shared.Cm(1),
                  shared.Cm(2.46),
                  shared.Cm(1),
                  shared.Cm(1.75),
                  shared.Cm(1),
                  shared.Cm(1.75),
                  shared.Cm(1.2),
                  shared.Cm(1.75),
                  shared.Cm(2),
                  )
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    def make_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'KEP'
                    for run in paragraph.runs:
                        run.font.bold = True

    def make_rows_font(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'KEP'

    def make_table_record(document):
        table = document.add_table(rows=1, cols=9)
        #table.style = 'table'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Merek'
        hdr_cells[2].text = 'Isi'
        #hdr_cells[3].text = ''
        hdr_cells[4].text = 'Jumlah'
        #hdr_cells[5].text = ''
        hdr_cells[6].text = 'Total'
        #hdr_cells[7].text = ''
        hdr_cells[8].text = 'HJE'

        # HORIZONTAL ALIGNMENT
        for i in range(9):
            cell = table.cell(0, i)
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # vertical aligment
        for i in range(9):
            table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # MAKE BOLD
        make_rows_bold(table.rows[0])

        # MERGE CELL
        a = table.cell(0, 2)
        b = table.cell(0, 3)
        c = table.cell(0, 4)
        d = table.cell(0, 5)
        e = table.cell(0, 6)
        f = table.cell(0, 7)

        A = a.merge(b)
        B = c.merge(d)
        B = e.merge(f)

        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        return table

    def make_table_header(document, jenis, jumlah_satuan_jenis, no_kep_bdn_full, tgl_kep_bdn):
        table = document.add_table(rows=9, cols=4)
        # ROW 1
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Menimbang'
        hdr_cells[1].text = ':'
        hdr_cells[2].text = 'a.'
        hdr_cells[3].text = 'Bahwa di Kantor Pengawasan dan Pelayanan Bea dan Cukai Tipe Madya Pabean C Tasikmalaya terdapat '+jumlah_satuan_jenis + \
            ' karena adanya pelanggaran terhadap Undang-Undang Nomor 11 Tahun 1995 tentang Cukai sebagaimana telah diubah dengan Undang-Undang Nomor 39 tahun 2007, sesuai dengan Barang Hasil Penindakan yang telah ditetapkan menjadi Barang yang Dikuasai Negara.'

        # ROW 2
        hdr_cells = table.rows[1].cells
        hdr_cells[2].text = 'b.'
        hdr_cells[3].text = 'Bahwa sehubungan dengan hal tersebut di atas, dipandang perlu untuk menetapkan barang-barang tersebut sebagai Barang Milik Negara.'

        # ROW 3
        hdr_cells = table.rows[2].cells

        # ROW 4
        hdr_cells = table.rows[3].cells
        hdr_cells[0].text = 'Mengingat'
        hdr_cells[1].text = ':'
        hdr_cells[2].text = '1.'
        hdr_cells[3].text = 'Undang-Undang Nomor 11 tahun 1995  tentang Cukai (Lembaran  Negara Republik Indonesia Tahun 1995 Nomor 76, Tambahan Lembaran Negara Republik Indonesia Nomor 3613) sebagaimana telah diubah dengan  Undang-Undang  Nomor 39 tahun 2007 (Lembaran Negara Republik Indonesia Tahun 2007 Nomor 105, Tambahan Lembaran Negara Republik Indonesia Nomor 4755);'

        # ROW 5
        hdr_cells = table.rows[4].cells

        hdr_cells[2].text = '2.'
        hdr_cells[3].text = 'Peraturan Menteri Keuangan Republik Indonesia Nomor: 39/PMK.04/2014 tanggal 19 Februari 2014 tentang Tata Cara Penyelesaian Barang Kena Cukai dan Barang - barang Lain yang Dirampas Untuk Negara atau yang Dikuasai Negara;'

        # ROW 6
        hdr_cells = table.rows[5].cells
        hdr_cells[2].text = '3.'
        hdr_cells[3].text = 'Keputusan Kepala Kantor Bea dan Cukai Tipe Madya Pabean C Tasikmalaya Nomor' + \
            no_kep_bdn_full+' tanggal '+tgl_kep_bdn+'.'

        # ROW 7
        hdr_cells = table.rows[6].cells

        # ROW 8
        hdr_cells = table.rows[7].cells
        hdr_cells[3].text = 'MEMUTUSKAN\n'
        table.cell(7, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ROW 9
        hdr_cells = table.rows[8].cells
        hdr_cells[0].text = 'Menetapkan'
        hdr_cells[1].text = ':'
        hdr_cells[3].text = 'KEPUTUSAN KEPALA KANTOR PENGAWASAN DAN PELAYANAN  BEA DAN CUKAI TIPE MADYA PABEAN C TASIKMALAYA TENTANG PENETAPAN BARANG KENA CUKAI BERUPA ' + \
            jenis.upper()+' SEBAGAI BARANG MILIK NEGARA'

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.style = 'KEP-1'
        widths = (shared.Cm(2.52),
                  shared.Cm(0.49),
                  shared.Cm(0.49),
                  shared.Cm(13.7),
                  )
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        return table

    # KEP BDN
    no_kep_bdn = "KEP-"+str(no_kep_bdn) + \
        "/WBC.09/KPP.MP.06/"+tgl_kep_bdn.strftime('%Y')
    tgl_kep_bdn = tgl_kep_bdn.strftime('%-d %B %Y')

    # KEP BMN
    no_kep_bmn = "KEP-"+str(no_kep_bmn) + \
        "/WBC.09/KPP.MP.06/"+tgl_kep_bmn.strftime('%Y')
    tgl_kep_bmn = tgl_kep_bmn.strftime('%-d %B %Y')

    # JUMLAH JENIS BARANG
    # jumlah = Barang.objects.filter(
    #     nomor_skep=id_skep).aggregate(Sum('jumlah_kemasan'))
    # jum = jumlah['jumlah_kemasan__sum']
    # jum_sat = all[0][7]
    # jenis = all[0][2]

    # jumlah_satuan_jenis = str(jum)+' '+str(jum_sat)+' '+str(jenis)

    # JUMLAH JENIS BARANG CARA BARU

    barang = Barang.objects.filter(nomor_skep=id_skep)

    def join_koma_dan(l):
        if not l:
            return ""
        elif len(l) == 1:
            return l[0]
        else:
            return ', '.join(l[:-1]) + " dan " + l[-1]

    barang_annotate = barang.values(jb=F('jenis_barang__jenis_barang'), sat=F('jenis_barang__jumlah_kemasan_satuan')).annotate(
        jumlah=Sum('jumlah_kemasan'))
    # JENIS BARANG TEXT
    list_barang = []
    for b in barang_annotate:
        text = str(b['jb'])
        list_barang.append(text)

    # FULL TEXT
    list_full = []
    for b in barang_annotate:
        text = str(b['jumlah'])+' ' + \
            str(b['sat'])+' '+str(b['jb'])
        list_full.append(text)

    jenis = join_koma_dan(list_barang)
    jumlah_satuan_jenis = join_koma_dan(list_full)

    # PENETAPAN
    lokasi = 'Tasikmalaya'
    kepala_kantor = 'Indriya Karyadi'
    nip = '19710905 199201 1 004'

    # SALINAN

    # SET DOCUMENT FORMAT
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(330)
    section.page_width = Mm(215)
    section.top_margin = Mm(20.0)
    section.bottom_margin = Mm(15.0)
    section.left_margin = Mm(25.0)
    section.right_margin = Mm(20.0)
    section.header_distance = Mm(10.0)
    section.footer_distance = Mm(12.7)

    # BASIC STYLE
    styles = document.styles
    style = styles.add_style('KEP', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    paragraph = style.paragraph_format
    font.name = 'Bookman Old Style'
    font.size = Pt(11)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.space_after = Pt(0)

    # BASIC STYLE AND JUSTIFY
    styles = document.styles
    style = styles.add_style('KEP-1', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    paragraph = style.paragraph_format
    font.name = 'Bookman Old Style'
    font.size = Pt(11)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.space_after = Pt(0)

    # STYLE INDENT AND JUSTIFY
    styles = document.styles
    style = styles.add_style('KEP-2', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    paragraph = style.paragraph_format
    font.name = 'Bookman Old Style'
    font.size = Pt(11)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.space_after = Pt(0)
    paragraph.left_indent = Mm(32.5)

    # STYLE INDENT AND CENTER
    styles = document.styles
    style = styles.add_style('KEP-2-Center', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    paragraph = style.paragraph_format
    font.name = 'Bookman Old Style'
    font.size = Pt(11)
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.space_after = Pt(0)
    paragraph.left_indent = Mm(32.5)

    para = document.add_paragraph(
        'KEMENTERIAN KEUANGAN REPUBLIK INDONESIA\nDIREKTORAT JENDERAL BEA DAN CUKAI\n\nKEPUTUSAN KEPALA KANTOR PENGAWASAN DAN PELAYANAN\nBEA DAN CUKAI TIPE MADYA PABEAN C TASIKMALAYA\n'+'NOMOR '+no_kep_bmn+'\n\nTENTANG\nPENETAPAN BARANG YANG DIKUASAI NEGARA\nMENJADI BARANG MILIK NEGARA\n\nKEPALA KANTOR PENGAWASAN DAN PELAYANAN BEA DAN CUKAI\nTIPE MADYA PABEAN C TASIKMALAYA\n\n', style='KEP')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #
    table_header = make_table_header(
        document, jenis, jumlah_satuan_jenis, no_kep_bdn, tgl_kep_bdn)

    para = document.add_paragraph('\nPasal 1', style='KEP-2-Center')
    para = document.add_paragraph('Barang-barang Kena Cukai berupa ' +
                                  jumlah_satuan_jenis+' dengan rincian:', style='KEP-2')

    table_record = make_table_record(document)

    i = 1
    for bdn, bmn, jb, mer, isi, isisat, jm, jmsat, hjecer, hjesat, nilai in all:
        total = isi * jm
        row_cells = table_record.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(mer)
        row_cells[2].text = str(isi)
        row_cells[3].text = str(isisat)
        row_cells[4].text = str(jm)
        row_cells[5].text = str(jmsat)
        row_cells[6].text = str(total)
        row_cells[7].text = str(isisat)
        row_cells[8].text = 'Rp'+str(hjecer)+'\n/'+str(jmsat)
        make_rows_font(table_record.rows[i])
        i = i+1

    set_col_widths_record(table_record)

    para = document.add_paragraph("adalah barang-barang hasil penindakan karena pelanggaran terhadap Undang-Undang Nomor 11 Tahun 1995 tentang Cukai sebagaimana telah diubah dengan Undang-Undang Nomor 39 tahun 2007,  tanpa dilekati pita cukai dan sesuai dengan Keputusan Kepala Kantor Pengawasan dan Pelayanan Bea dan Cukai Tipe Madya Pabean C Tasikmalaya Nomor " + no_kep_bdn + " tanggal "+tgl_kep_bdn+".", style='KEP-2')

    para = document.add_paragraph('\nPasal 2', style='KEP-2-Center')
    para = document.add_paragraph(
        'Barang-barang sebagaimana dimaksud dalam Pasal 1 diatas sesuai dengan Pasal 54 Undang-Undang Nomor 11 Tahun 1995 tentang Cukai sebagaimana telah diubah dengan Undang-Undang Nomor 39 Tahun 2007, maka ditetapkan sebagai Barang Milik Negara.', style='KEP-2')

    para = document.add_paragraph('\nPasal 3', style='KEP-2-Center')
    para = document.add_paragraph(
        'Barang yang Dikuasai Negara sebagaimana dimaksud dalam Pasal 1 di atas ditimbun di Gudang Tempat Penimbunan Barang Hasil Penindakan di bawah Kantor Pengawasan dan Pelayanan Bea dan Cukai Tipe Madya Pabean C Tasikmalaya.', style='KEP-2')

    para = document.add_paragraph('\nPasal 4', style='KEP-2-Center')
    para = document.add_paragraph(
        'Penyelesaian lebih lanjut atas Barang Milik Negara sebagaimana dimaksud dalam pasal 2 di atas, dilakukan pemusnahan oleh Pejabat Bea dan Cukai.', style='KEP-2')

    para = document.add_paragraph('\nPasal 5', style='KEP-2-Center')
    para = document.add_paragraph(
        'Keputusan ini mulai berlaku pada tanggal ditetapkan, dengan ketentuan apabila di kemudian hari terdapat kekeliruan akan diadakan perbaikan sebagaimana mestinya.', style='KEP-2')

    para = document.add_paragraph(
        '\nSalinan Keputusan Kepala Kantor Pengawasan dan Pelayanan Bea dan Cukai Tipe Madya Pabean C Tasikmalaya ini disampaikan kepada\n1. Direktur Teknis dan Fasilitas Cukai;\n2. Direktur Penindakan dan Penyidikan;\n3. Kepala Kantor Wilayah DJBC Jawa Barat.\n\n', style='KEP')
    para.paragraph_format.left_indent = Mm(32.5)

    para = document.add_paragraph(
        'Ditetapkan di : '+lokasi+'\nPada tanggal : '+tgl_kep_bdn+'\n\nKepala Kantor,\n\n\n\n'+kepala_kantor+'\nNIP '+nip, style='KEP')
    para.paragraph_format.left_indent = Mm(100.0)

    return document
